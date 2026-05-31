from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("reports/kaixiang_individual_report")
OUT_DOCX = OUT_DIR / "Kaixiang_Sheng_Individual_Report_Text_Tables.docx"


def set_run_font(run, size=9, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_para_font(paragraph, size=9):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=run.bold, italic=run.italic)


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7 if level == 1 else 4)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, size=12 if level == 1 else 10, bold=True, color=(31, 78, 121) if level == 1 else None)
    return paragraph


def add_para(document, text, size=8.8):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    set_para_font(paragraph, size=size)
    return paragraph


def add_bullet(document, text, size=8.5):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.add_run(text)
    set_para_font(paragraph, size=size)
    return paragraph


def add_figure_placeholder(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, size=8, italic=True, color=(89, 89, 89))
    return paragraph


def compact_table(table, font_size=7.4):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, bold=(row_index == 0))
                if row_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(document, headers, rows, font_size=7.4):
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    compact_table(table, font_size=font_size)
    return table


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Individual Report - Kaixiang Sheng")
    set_run_font(title_run, size=15, bold=True, color=(31, 78, 121))

    subtitle = document.add_paragraph("Facial Recognition with Emotion and Liveness")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_font(subtitle, size=8.5)

    add_heading(document, "1. Individual Contributions", level=1)
    add_para(
        document,
        "My individual work focused on three parts of the attendance system: metric-learning face verification, "
        "anti-spoofing / liveness detection, and glasses detection as an independent innovative feature. I built "
        "training scripts, evaluation scripts, model comparison outputs, and integration wrappers so that each "
        "module could be tested quantitatively and through webcam-based behaviour.",
    )
    add_bullet(document, "Metric recognition: trained and compared four Kaixiang metric-learning models and evaluated Euclidean distance and cosine similarity.")
    add_bullet(document, "Liveness detection: trained two anti-spoofing models and tuned the final integration behaviour with expanded crops and thresholding.")
    add_bullet(document, "Innovation: implemented glasses detection as a separate binary module that reports whether a registered user is wearing glasses.")

    add_heading(document, "2. Datasets and Evaluation Protocol", level=1)
    add_para(
        document,
        "The original 11-785 Face Verification dataset was first tested, but smoke-test performance was not stable "
        "enough for the final metric-recognition pipeline. Since the specification allows additional public data, "
        "I used larger task-specific datasets and kept final evaluations on fixed local test splits.",
    )
    add_bullet(document, "Original dataset reference: https://www.kaggle.com/c/11-785-fall-20-homework-2-part-2/overview/evaluation")
    add_bullet(document, "Metric recognition dataset: VGGFace2, https://www.kaggle.com/datasets/hearfool/vggface2. Local split: 480 training identities, 30 validation identities, and 30 test identities. Final evaluation used datasets/recognition/test with 1200 verification pairs.")
    add_bullet(document, "Liveness dataset: LCC-FASD, https://www.kaggle.com/datasets/faber24/lcc-fasd. The local structure used real/spoof folders with an approximately 5:1:1 train/validation/test split. Final test set: 200 images, 100 real and 100 spoof.")
    add_bullet(document, "Glasses dataset: Face Cropped Glasses vs No Glasses, https://www.kaggle.com/datasets/sehriyarmemmedli/facecropped-glasses-vs-noglasses-dataset. Final test set: 3376 images.")

    add_heading(document, "3. Methods", level=1)
    add_para(
        document,
        "All final models used transfer learning with task-specific training, rather than directly using fully "
        "pre-trained models. The common input format was a cropped RGB face image resized to 224 x 224 x 3. "
        "For recognition, each face was mapped to a compact L2-normalised embedding and pairs were compared "
        "using Euclidean distance and cosine similarity. ROC-AUC was the main metric because it evaluates pair "
        "ranking across thresholds.",
    )
    add_bullet(document, "Contrastive metric models: Siamese ResNet18 and Siamese MobileNetV2, contrastive loss, batch size 16, 5 head epochs + 15 fine-tuning epochs, 3000 pairs per epoch, head LR 1e-3, fine-tune LR 1e-4.")
    add_bullet(document, "Triplet metric models: Triplet ResNet18 and Triplet MobileNetV2 with batch-hard triplet training. The sampler used P=8 identities and K=4 images per identity. Triplet ResNet18 used margin 0.3; Triplet MobileNetV2 used margin 0.2 in the final run.")
    add_bullet(document, "Liveness models: MobileNetV2 and EfficientNetB0 with custom binary heads, binary cross-entropy objective, 5 head epochs + 15 fine-tuning epochs, early stopping, and final-backbone-block fine-tuning.")
    add_bullet(document, "Glasses models: MobileNetV2 and EfficientNetB0 with binary heads, BCEWithLogitsLoss, class weighting, 3 head epochs + 8 fine-tuning epochs, and capped training samples per class for practical runtime.")

    add_heading(document, "4. Results and Model Selection", level=1)
    add_heading(document, "4.1 Metric Learning Recognition", level=2)
    metric_rows = [
        ["Contrastive Siamese + ResNet18", "Contrastive", "0.8375", "0.7708", "0.7703", "0.8375", "15.55"],
        ["Contrastive Siamese + MobileNetV2", "Contrastive", "0.7739", "0.7142", "0.7491", "0.7739", "29.33"],
        ["Triplet + ResNet18", "Triplet", "0.8880", "0.8092", "0.8044", "0.8883", "16.75"],
        ["Triplet + MobileNetV2", "Triplet", "0.7794", "0.7192", "0.7310", "0.7791", "29.82"],
    ]
    add_table(document, ["Model", "Loss", "Euc. AUC", "Euc. Acc.", "Euc. F1", "Cos. AUC", "FPS"], metric_rows)
    add_para(
        document,
        "Triplet + ResNet18 was selected as my final recognition model because it achieved the highest ROC-AUC "
        "and accuracy on the shared recognition test pairs. Euclidean and cosine results were very close because "
        "the embeddings were L2-normalised; the deployed wrapper used Euclidean thresholding for verification.",
        size=8.5,
    )
    add_figure_placeholder(document, "Insert Figure 1 here: ROC curves for the four Kaixiang metric models.")
    add_figure_placeholder(document, "Insert Figure 2 here: positive vs negative score distribution for Triplet + ResNet18.")

    add_heading(document, "4.2 Anti-Spoofing / Liveness Detection", level=2)
    liveness_rows = [
        ["MobileNetV2 + binary head", "0.9447", "0.8750", "0.8756", "0.8713", "0.8800", "37.62"],
        ["EfficientNetB0 + binary head", "0.9724", "0.9150", "0.9128", "0.9368", "0.8900", "30.92"],
    ]
    add_table(document, ["Model", "AUC", "Acc.", "F1", "Precision", "Recall", "FPS"], liveness_rows)
    add_para(
        document,
        "EfficientNetB0 + binary head was selected for liveness because it produced the strongest ROC-AUC, "
        "accuracy, and F1-score. MobileNetV2 was faster, but liveness is security-critical, so reliability was "
        "prioritised over the small speed advantage.",
        size=8.5,
    )
    add_figure_placeholder(document, "Insert Figure 3 here: liveness ROC curve or liveness score distribution for EfficientNetB0.")

    add_heading(document, "4.3 Innovative Feature: Glasses Detection", level=2)
    add_para(
        document,
        "Glasses detection was designed as an independent D/HD extension, separate from recognition, liveness, "
        "and emotion detection. The rationale was to add useful appearance metadata for an attendance system, "
        "for example whether the registered person is currently wearing glasses. This can make the final demo "
        "more informative without changing the core identity decision.",
        size=8.5,
    )
    glasses_rows = [
        ["MobileNetV2 + binary head", "0.9932", "0.9892", "0.9942", "0.9917", "119.12"],
        ["EfficientNetB0 + binary head", "0.9923", "0.9884", "0.9927", "0.9906", "103.52"],
    ]
    add_table(document, ["Model", "Acc.", "Precision", "Recall", "F1", "FPS"], glasses_rows)
    add_para(
        document,
        "Both glasses models performed strongly, but MobileNetV2 was slightly more accurate and faster. Therefore, "
        "MobileNetV2 + binary head was selected as the final glasses detection model.",
        size=8.5,
    )
    add_figure_placeholder(document, "Insert Figure 4 here: glasses model comparison bar chart.")

    add_heading(document, "5. Challenges and Resolutions", level=1)
    add_bullet(document, "Metric model instability: webcam testing showed that some similar-looking identities had very small embedding gaps. I addressed this with better triplet training, identity-disjoint testing, thresholding, ambiguity rejection, and gallery-based matching.")
    add_bullet(document, "Dataset choice: the original verification dataset was not stable enough for my model development, so I used VGGFace2 for metric learning and documented the final fixed test-pair protocol.")
    add_bullet(document, "Liveness false accepts: tight crops sometimes ignored spoof context. I improved integration testing with expanded liveness crops and adjusted the decision threshold.")
    add_bullet(document, "Runtime constraints: training and evaluation ran mostly on CPU, so I used two-stage fine-tuning, early stopping where appropriate, and capped samples for glasses training.")

    add_heading(document, "6. Reflection", level=1)
    add_para(
        document,
        "My final selected models were Triplet + ResNet18 for metric recognition, EfficientNetB0 + binary head for "
        "liveness detection, and MobileNetV2 + binary head for glasses detection. The biggest lesson was that strong "
        "offline metrics are necessary but not sufficient: webcam behaviour also depends on crop quality, lighting, "
        "face pose, and registered gallery quality. Combining quantitative evaluation with real-time testing made the "
        "final modules more practical for the integrated attendance system.",
    )

    document.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
