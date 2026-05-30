from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUT = Path("reports/kaixiang_zhongyu_recomparison/metric_liveness_group_report_section.docx")


recognition_rows = [
    ["Kaixiang", "Contrastive Siamese + ResNet18", "Contrastive loss", "0.8375", "0.7708", "0.7703", "0.8375", "15.55"],
    ["Kaixiang", "Contrastive Siamese + MobileNetV2", "Contrastive loss", "0.7739", "0.7142", "0.7491", "0.7739", "29.33"],
    ["Kaixiang", "Triplet + ResNet18", "Triplet loss", "0.8880", "0.8092", "0.8044", "0.8883", "16.75"],
    ["Kaixiang", "Triplet + MobileNetV2", "Triplet loss", "0.7794", "0.7192", "0.7310", "0.7791", "29.82"],
    ["Zhongyu", "Triplet + ResNet34", "Triplet loss", "0.7720", "0.7033", "0.7206", "0.7720", "7.68"],
    ["Zhongyu", "Triplet + EfficientNet-B0", "Triplet loss", "0.5992", "0.5983", "0.5170", "0.5992", "21.26"],
]

liveness_rows = [
    ["Kaixiang", "MobileNetV2 + binary head", "Binary cross-entropy", "0.9447", "0.8750", "0.8756", "0.8713", "0.8800", "37.62"],
    ["Kaixiang", "EfficientNetB0 + binary head", "Binary cross-entropy", "0.9724", "0.9150", "0.9128", "0.9368", "0.8900", "30.92"],
    ["Zhongyu", "DenseNet121 + binary head", "Binary cross-entropy", "0.9053", "0.8800", "0.8737", "0.9222", "0.8300", "10.70"],
    ["Zhongyu", "ResNet50V2 + binary head", "Binary cross-entropy", "0.8813", "0.8200", "0.8235", "0.8077", "0.8400", "14.48"],
]


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        header_cells[i].text = text
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    return table


def add_caption(document, text):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(2)


def main():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Face Verification and Anti-Spoofing Report Section")
    title_run.bold = True
    title_run.font.size = Pt(16)

    document.add_paragraph(
        "This section is written as a concise group-report draft for the face verification "
        "and anti-spoofing components. Detailed implementation code and individual tuning "
        "rationale can be kept in the individual reports."
    )

    document.add_heading("Dataset and Evaluation Protocol", level=2)
    document.add_paragraph(
        "The original project data was extended with public Kaggle datasets because the supplied "
        "verification data gave limited generalisation in webcam testing. For metric-learning "
        "recognition, Kaixiang used VGGFace2 data, organised into identity-disjoint train, "
        "validation and test folders. The local split used for the final comparison contains "
        "176,398 training images from 480 identities, 11,083 validation images from 30 identities, "
        "and 10,212 test images from 30 identities. Zhongyu trained additional metric-learning "
        "models using CelebA-500 label folders with an approximate 5:1:1 train/validation/test "
        "split. To make the final comparison fair, all recognition models were evaluated on the "
        "same shared test protocol: datasets/recognition/test, sampled into 1,200 deterministic "
        "verification pairs."
    )
    document.add_paragraph(
        "For liveness detection, the data was organised as a binary real/spoof dataset under "
        "datasets/liveness/train, val and test. The final local split follows a 5:1:1 ratio per "
        "class: 500 real and 500 spoof images for training, 100 real and 100 spoof images for "
        "validation, and 100 real and 100 spoof images for testing. The final anti-spoofing "
        "comparison used datasets/liveness/test, so every model was evaluated on the same 200 "
        "images. Note: before final submission, verify the exact Kaggle URL for the liveness "
        "source in the reference list; the current local folder is a real/spoof anti-spoofing set."
    )

    document.add_heading("Face Verification: Methodology", level=2)
    document.add_paragraph(
        "The face verification module was designed as an open-set metric-learning task. Each model "
        "maps a 224 x 224 RGB face image into a compact L2-normalised embedding. Verification is "
        "performed by comparing two embeddings and applying a threshold to decide whether the pair "
        "belongs to the same identity. Six metric-learning models were compared: four Kaixiang "
        "models and two Zhongyu models. Kaixiang compared contrastive Siamese models and triplet "
        "loss models using ResNet18 and MobileNetV2 backbones. Zhongyu trained triplet loss models "
        "using ResNet34 and EfficientNet-B0. Transfer learning was used, but the models were not "
        "used as fully pre-trained feature extractors: the new embedding heads were trained and "
        "selected backbone blocks were fine-tuned."
    )
    document.add_paragraph(
        "The main evaluation metrics were ROC-AUC, accuracy, F1 score and inference speed. Both "
        "Euclidean distance and cosine similarity were evaluated as required by the specification. "
        "For Euclidean distance, lower distance indicates higher similarity; for cosine similarity, "
        "higher score indicates higher similarity."
    )

    document.add_heading("Face Verification: Results", level=2)
    add_table(
        document,
        ["Owner", "Model", "Training loss", "Euc. AUC", "Euc. Acc.", "Euc. F1", "Cos. AUC", "FPS"],
        recognition_rows,
    )
    add_caption(
        document,
        "Insert Figure R1 here: reports/kaixiang_zhongyu_recomparison/recognition_model_auc_accuracy_fps.png"
    )
    add_caption(
        document,
        "If space allows, insert Figure R2: recognition_roc_curves.png. Optional for section 2.3: "
        "recognition_best_model_distance_metric_comparison.png and recognition_best_model_score_distribution.png."
    )

    document.add_heading("Face Verification: Discussion", level=2)
    document.add_paragraph(
        "The best recognition result was obtained by Kaixiang's Triplet + ResNet18 model, with the "
        "highest Euclidean ROC-AUC (0.8880), cosine ROC-AUC (0.8883), accuracy (0.8092 using "
        "Euclidean distance) and F1 score (0.8044). This suggests that direct triplet supervision "
        "created a more discriminative embedding space than the earlier contrastive Siamese models. "
        "Although MobileNetV2 was faster, its accuracy and AUC were lower, so ResNet18 was selected "
        "as the final face verification model. Euclidean distance and cosine similarity produced "
        "very similar AUC values because the embeddings were L2-normalised; Euclidean distance was "
        "kept as the primary thresholding metric for integration, while cosine similarity was "
        "reported for comparison."
    )

    document.add_heading("Anti-Spoofing: Methodology", level=2)
    document.add_paragraph(
        "The anti-spoofing module was trained as a binary image classification task. Each cropped "
        "face image is classified as REAL or SPOOF. Four models were compared: Kaixiang's "
        "MobileNetV2 + binary head and EfficientNetB0 + binary head, and Zhongyu's DenseNet121 "
        "and ResNet50V2 binary classifiers. All models used transfer learning with a binary "
        "classification head and binary cross-entropy loss. The training scheme followed the "
        "project rule that fully pre-trained models cannot be used directly: the classifier heads "
        "were trained and the models were further tuned rather than used as fixed pre-trained models."
    )
    document.add_paragraph(
        "The main anti-spoofing metrics were ROC-AUC, accuracy, precision, recall, F1 score and "
        "FPS. ROC-AUC was treated as the primary selection criterion because it measures ranking "
        "quality across thresholds, while accuracy and F1 describe performance at the selected "
        "decision threshold."
    )

    document.add_heading("Anti-Spoofing: Results", level=2)
    add_table(
        document,
        ["Owner", "Model", "Loss", "AUC", "Acc.", "F1", "Precision", "Recall", "FPS"],
        liveness_rows,
    )
    add_caption(
        document,
        "Insert Figure L1 here: reports/kaixiang_zhongyu_recomparison/liveness_model_auc_accuracy_fps.png"
    )
    add_caption(
        document,
        "Optional supporting figures: liveness_roc_curves.png and "
        "liveness_best_model_score_distribution.png."
    )

    document.add_heading("Anti-Spoofing: Discussion", level=2)
    document.add_paragraph(
        "Kaixiang's EfficientNetB0 + binary head achieved the strongest overall anti-spoofing "
        "performance, with the highest ROC-AUC (0.9724), accuracy (0.9150), and F1 score (0.9128). "
        "It also reduced false positives compared with the other models, which is important for an "
        "attendance system because spoofed faces should be rejected reliably. MobileNetV2 was the "
        "fastest Kaixiang liveness model, but EfficientNetB0 provided a better balance of accuracy "
        "and robustness. Therefore, EfficientNetB0 was selected as the final liveness model for "
        "system integration."
    )

    document.add_heading("Final Selection for Integration", level=2)
    add_bullet(
        document,
        "Face verification: Kaixiang Triplet Loss Metric Learning + ResNet18 embedding network."
    )
    add_bullet(
        document,
        "Anti-spoofing: Kaixiang EfficientNetB0 + binary classification head."
    )
    document.add_paragraph(
        "These models were selected because they gave the best ROC-AUC and overall performance in "
        "the shared final comparison while remaining compatible with the webcam-based attendance "
        "pipeline."
    )

    document.add_heading("Recommended Figure Use", level=2)
    add_bullet(document, "Use only 2-3 figures in the group report if page space is tight.")
    add_bullet(document, "Recommended minimum: recognition_model_auc_accuracy_fps.png and liveness_model_auc_accuracy_fps.png.")
    add_bullet(document, "Add recognition_roc_curves.png if the report needs explicit coverage of ROC/AUC.")
    add_bullet(document, "Add recognition_best_model_distance_metric_comparison.png if the report needs explicit coverage of Euclidean vs cosine distance.")
    add_bullet(document, "Keep score-distribution plots for appendix, individual report, or oral explanation unless there is enough space.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
