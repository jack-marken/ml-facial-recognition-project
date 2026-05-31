from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUT_DIR = Path("reports/kaixiang_individual_report")
OUT_DOCX = OUT_DIR / "Kaixiang_Sheng_Individual_Report_Draft.docx"


def set_font(paragraph, size=9, bold=False, italic=False):
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(13 if level == 1 else 10.5)
    paragraph.paragraph_format.space_before = Pt(6 if level == 1 else 4)
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def add_para(document, text, size=9):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    set_font(paragraph, size=size)
    return paragraph


def add_bullet(document, text, size=8.8):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0
    set_font(paragraph, size=size)
    return paragraph


def add_caption(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    set_font(paragraph, size=7.5, italic=True)
    return paragraph


def compact_table(table, font_size=7.2):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    compact_table(table)
    return table


def add_image(document, path, width):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_two_images(document, left_path, right_path, left_caption, right_caption):
    table = document.add_table(rows=2, cols=2)
    table.autofit = True
    for col, path in enumerate([left_path, right_path]):
        paragraph = table.rows[0].cells[col].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(path), width=Inches(3.25))
    for col, caption in enumerate([left_caption, right_caption]):
        paragraph = table.rows[1].cells[col].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = caption
        set_font(paragraph, size=7.2, italic=True)
    compact_table(table, font_size=7.2)
    for row in table.rows:
        for cell in row.cells:
            cell._tc.get_or_add_tcPr()
    return table


def main():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Individual Report - Kaixiang Sheng")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(16)
    subtitle = document.add_paragraph("Facial Recognition with Emotion and Liveness")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle, size=9)

    add_heading(document, "1. Individual Contributions", level=1)
    add_para(
        document,
        "My main contributions were metric-learning face verification, anti-spoofing / liveness detection, "
        "and glasses detection as an individual innovative feature. I implemented training, evaluation, "
        "comparison, and integration wrappers for these modules, then selected the final models based on "
        "shared test-set performance and real-time webcam behaviour.",
    )
    add_bullet(document, "Metric learning recognition: trained and compared four Kaixiang models, built gallery-based recognition, and tested Euclidean and cosine verification.")
    add_bullet(document, "Anti-spoofing: trained MobileNetV2 and EfficientNetB0 binary liveness classifiers and tuned webcam integration thresholds.")
    add_bullet(document, "Innovation: developed an independent glasses detection module for additional user appearance metadata.")

    add_heading(document, "2. Datasets and Evaluation Protocol", level=1)
    add_para(
        document,
        "The supplied 11-785 Face Verification dataset was initially tested, but smoke tests were not stable "
        "enough for robust metric recognition. The project specification allows additional data sources, so "
        "I used public datasets and kept identity-disjoint or class-balanced local splits.",
    )
    add_bullet(document, "Metric recognition dataset: VGGFace2, https://www.kaggle.com/datasets/hearfool/vggface2. Local split: 480 training identities, 30 validation identities, and 30 test identities. Final evaluation used datasets/recognition/test with 1200 verification pairs.")
    add_bullet(document, "Liveness dataset: LCC-FASD, https://www.kaggle.com/datasets/faber24/lcc-fasd. It was arranged into real/spoof folders with an approximately 5:1:1 train/validation/test split. Final evaluation used 200 test images: 100 real and 100 spoof.")
    add_bullet(document, "Glasses dataset: local datasets/glasses split into with_glasses and without_glasses. Final test set contained 3376 images.")

    add_heading(document, "3. Key Methods", level=1)
    add_para(
        document,
        "All models used transfer learning with further task-specific training, not fully pre-trained models directly. "
        "For metric recognition, I compared contrastive Siamese models and batch-hard triplet models. The models "
        "used 224 x 224 RGB crops, ImageNet normalization, L2-normalized embeddings, and threshold-based verification.",
    )
    add_bullet(document, "Contrastive models: Siamese ResNet18 and Siamese MobileNetV2, contrastive loss, batch size 16, 5 head epochs + 15 fine-tuning epochs, 3000 pairs per epoch, head LR 1e-3, fine-tune LR 1e-4.")
    add_bullet(document, "Triplet models: Triplet ResNet18 and Triplet MobileNetV2 with P=8 identities and K=4 images per identity. Triplet ResNet18 used margin 0.3; Triplet MobileNetV2 used margin 0.2 for the final run.")
    add_bullet(document, "Liveness models: MobileNetV2 and EfficientNetB0 with custom binary heads, binary cross-entropy, 5 head epochs + 15 fine-tuning epochs, early stopping, and final-block fine-tuning.")
    add_bullet(document, "Glasses models: MobileNetV2 and EfficientNetB0 with binary heads, BCEWithLogitsLoss, class weighting, 3 head epochs + 8 fine-tuning epochs, and limited training samples per class for runtime control.")

    add_heading(document, "4. Results", level=1)
    add_heading(document, "4.1 Metric Learning Recognition", level=2)
    metric_rows = [
        ["Contrastive ResNet18", "Contrastive", "0.8375", "0.7708", "0.7703", "0.8375", "15.55"],
        ["Contrastive MobileNetV2", "Contrastive", "0.7739", "0.7142", "0.7491", "0.7739", "29.33"],
        ["Triplet ResNet18", "Triplet", "0.8880", "0.8092", "0.8044", "0.8883", "16.75"],
        ["Triplet MobileNetV2", "Triplet", "0.7794", "0.7192", "0.7310", "0.7791", "29.82"],
    ]
    add_table(document, ["Model", "Loss", "Euc AUC", "Euc Acc", "Euc F1", "Cos AUC", "FPS"], metric_rows)
    add_para(document, "Triplet ResNet18 achieved the best recognition result and was selected for final deployment.", size=8.6)
    add_two_images(
        document,
        OUT_DIR / "metric_kaixiang_model_comparison.png",
        OUT_DIR / "metric_kaixiang_roc_curves.png",
        "Figure 1. Kaixiang metric model comparison.",
        "Figure 2. ROC curves for my four metric models.",
    )
    add_two_images(
        document,
        OUT_DIR / "metric_best_distance_metric_comparison.png",
        OUT_DIR / "metric_best_score_distribution.png",
        "Figure 3. Euclidean vs cosine on the best model.",
        "Figure 4. Positive/negative score distribution.",
    )

    add_heading(document, "4.2 Anti-Spoofing / Liveness Detection", level=2)
    liveness_rows = [
        ["MobileNetV2 + binary head", "0.9447", "0.8750", "0.8756", "0.8713", "0.8800", "37.62"],
        ["EfficientNetB0 + binary head", "0.9724", "0.9150", "0.9128", "0.9368", "0.8900", "30.92"],
    ]
    add_table(document, ["Model", "AUC", "Acc", "F1", "Precision", "Recall", "FPS"], liveness_rows)
    add_para(document, "EfficientNetB0 was selected because it achieved the strongest ROC-AUC, accuracy, and F1-score.", size=8.6)
    add_two_images(
        document,
        OUT_DIR / "liveness_kaixiang_model_comparison.png",
        OUT_DIR / "liveness_best_score_distribution.png",
        "Figure 5. Kaixiang liveness model comparison.",
        "Figure 6. EfficientNetB0 score distribution.",
    )

    add_heading(document, "4.3 Innovative Feature: Glasses Detection", level=2)
    add_para(
        document,
        "Glasses detection was added as an independent innovative feature. It predicts whether the detected user "
        "is wearing glasses and can provide extra appearance information during attendance. The feature is separate "
        "from recognition and liveness detection.",
        size=8.7,
    )
    glasses_rows = [
        ["MobileNetV2 + binary head", "0.9932", "0.9892", "0.9942", "0.9917", "119.12"],
        ["EfficientNetB0 + binary head", "0.9923", "0.9884", "0.9927", "0.9906", "103.52"],
    ]
    add_table(document, ["Model", "Acc", "Precision", "Recall", "F1", "FPS"], glasses_rows)
    add_two_images(
        document,
        OUT_DIR / "glasses_kaixiang_model_comparison.png",
        OUT_DIR / "glasses_mobilenetv2_confusion_matrix.png",
        "Figure 7. Glasses model comparison.",
        "Figure 8. Final glasses confusion matrix.",
    )

    add_heading(document, "5. Challenges and Resolutions", level=1)
    add_bullet(document, "Dataset stability: the original verification dataset did not generalise well in smoke tests, so I used VGGFace2 and identity-disjoint train/validation/test splits.")
    add_bullet(document, "Model compliance: to satisfy the rule against fully pre-trained models, all final models replaced the classification head and were further trained or fine-tuned.")
    add_bullet(document, "Webcam recognition instability: similar-looking identities caused small embedding gaps, so I used thresholding, ambiguity rejection, gallery embeddings, and temporal smoothing in integration testing.")
    add_bullet(document, "Liveness false acceptance: tight face crops sometimes missed spoof context, so I tested expanded liveness crops and adjusted the decision threshold for better real-world behaviour.")

    add_heading(document, "6. Final Reflection", level=1)
    add_para(
        document,
        "Overall, Triplet ResNet18 was the strongest metric-learning model, EfficientNetB0 was the strongest liveness "
        "model, and MobileNetV2 was the best glasses model. The main lesson was that good offline AUC does not always "
        "guarantee stable webcam behaviour, so final integration required both quantitative evaluation and real-time "
        "testing under practical lighting, crop, and identity-similarity conditions.",
    )

    document.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
