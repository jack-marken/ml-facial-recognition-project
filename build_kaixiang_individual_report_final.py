from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("reports/kaixiang_individual_report")
OUT_DOCX = OUT_DIR / "Kaixiang_Sheng_Individual_Report_Final.docx"


BLUE = (31, 78, 121)
GREY = (89, 89, 89)


def set_run(run, size=8.6, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_runs(paragraph, size=8.6):
    for run in paragraph.runs:
        set_run(run, size=size, bold=run.bold, italic=run.italic)


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_run(run, size=12 if level == 1 else 9.5, bold=True, color=BLUE if level == 1 else None)
    return paragraph


def add_para(document, text, size=8.5):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(2.5)
    paragraph.paragraph_format.line_spacing = 1.0
    set_paragraph_runs(paragraph, size=size)
    return paragraph


def add_bullet(document, text, size=8.25):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.add_run(text)
    set_paragraph_runs(paragraph, size=size)
    return paragraph


def add_figure_note(document, figure_no, image_name, purpose):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(2.5)
    run = paragraph.add_run(f"Figure {figure_no}: insert {image_name} here - {purpose}.")
    set_run(run, size=7.6, italic=True, color=GREY)
    return paragraph


def compact_table(table, font_size=7.1):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                if row_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run(run, size=font_size, bold=(row_index == 0))


def add_table(document, headers, rows, font_size=7.1):
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    compact_table(table, font_size=font_size)
    return table


def add_key_value_table(document, rows):
    table = document.add_table(rows=0, cols=2)
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    compact_table(table, font_size=7.3)
    return table


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Individual Report - Kaixiang Sheng")
    set_run(title_run, size=15, bold=True, color=BLUE)

    subtitle = document.add_paragraph("Facial Recognition with Emotion and Liveness")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_runs(subtitle, size=8.4)

    add_heading(document, "1. Contributions", level=1)
    add_para(
        document,
        "My individual contribution covered three modules: metric-learning face verification, anti-spoofing / "
        "liveness detection, and glasses detection as my innovative feature. I implemented model training, "
        "evaluation, comparison, checkpoint selection, and integration wrappers so the modules could be tested "
        "both quantitatively and in webcam-based scenarios.",
    )
    add_bullet(document, "Metric recognition: four Kaixiang models were trained and compared, including contrastive Siamese and triplet-loss embedding networks.")
    add_bullet(document, "Liveness detection: MobileNetV2 and EfficientNetB0 binary anti-spoofing models were trained and tuned for real/spoof prediction.")
    add_bullet(document, "Innovation: glasses detection was implemented as an independent attendance-system feature, separate from recognition and liveness.")

    add_heading(document, "2. Datasets and Test Protocol", level=1)
    add_para(
        document,
        "The original 11-785 Face Verification dataset was tested first, but the smoke-test results were not stable "
        "enough for my final metric-recognition models. The project specification allows additional public datasets, "
        "so I used task-specific Kaggle datasets and fixed local test splits for fair comparison.",
    )
    add_key_value_table(
        document,
        [
            (
                "Original dataset",
                "11-785 Fall 20 Homework 2 Part 2: https://www.kaggle.com/c/11-785-fall-20-homework-2-part-2/overview/evaluation",
            ),
            (
                "Metric dataset",
                "VGGFace2: https://www.kaggle.com/datasets/hearfool/vggface2. Local split: 480 train identities, 30 validation identities, 30 test identities. Final test: datasets/recognition/test with 1200 verification pairs.",
            ),
            (
                "Liveness dataset",
                "LCC-FASD: https://www.kaggle.com/datasets/faber24/lcc-fasd. Local split was approximately train:val:test = 5:1:1. Final test: datasets/liveness/test, 200 images total, 100 real and 100 spoof.",
            ),
            (
                "Glasses dataset",
                "Face Cropped Glasses vs No Glasses: https://www.kaggle.com/datasets/sehriyarmemmedli/facecropped-glasses-vs-noglasses-dataset. Final test: 3376 cropped face images.",
            ),
        ],
    )

    add_heading(document, "3. Methods", level=1)
    add_para(
        document,
        "All models used transfer learning with further task-specific training instead of directly deploying fully "
        "pre-trained models. Inputs were cropped RGB face images resized to 224 x 224 x 3. For metric recognition, "
        "the CNN output was converted into an L2-normalised embedding, and face pairs were compared using Euclidean "
        "distance and cosine similarity. ROC-AUC was used as the main verification metric because it evaluates pair "
        "ranking across possible thresholds.",
    )
    add_bullet(document, "Contrastive metric models: Siamese ResNet18 and Siamese MobileNetV2, contrastive loss, batch size 16, 5 head epochs + 15 fine-tuning epochs, 3000 pairs per epoch, head LR 1e-3, fine-tune LR 1e-4, early stopping patience 5.")
    add_bullet(document, "Triplet metric models: Triplet ResNet18 and Triplet MobileNetV2 with batch-hard triplet training. The sampler used P=8 identities and K=4 images per identity. Triplet ResNet18 used margin 0.3; Triplet MobileNetV2 used margin 0.2 in the final run.")
    add_bullet(document, "Liveness models: MobileNetV2 and EfficientNetB0 with custom binary heads, binary cross-entropy, 5 head epochs + 15 fine-tuning epochs, final-block fine-tuning, and early stopping patience 5.")
    add_bullet(document, "Glasses models: MobileNetV2 and EfficientNetB0 with binary heads, BCEWithLogitsLoss, positive class weighting, 3 head epochs + 8 fine-tuning epochs, capped training samples per class, and early stopping patience 4.")

    add_heading(document, "4. Results and Model Selection", level=1)
    add_heading(document, "4.1 Metric Learning Recognition", level=2)
    metric_rows = [
        ["Siamese ResNet18", "Contrastive", "0.8375", "0.7708", "0.7703", "0.8375", "24.19"],
        ["Siamese MobileNetV2", "Contrastive", "0.7739", "0.7142", "0.7491", "0.7739", "27.19"],
        ["Triplet ResNet18", "Triplet", "0.8880", "0.8092", "0.8044", "0.8883", "12.86"],
        ["Triplet MobileNetV2", "Triplet", "0.7794", "0.7192", "0.7310", "0.7791", "33.05"],
    ]
    add_table(document, ["Model", "Loss", "Euc. AUC", "Euc. Acc.", "Euc. F1", "Cos. AUC", "FPS"], metric_rows)
    add_para(
        document,
        "Triplet ResNet18 was selected as my final recognition model because it achieved the strongest ROC-AUC and "
        "accuracy on the recognition test pairs. Euclidean distance and cosine similarity produced very similar "
        "rankings because the embeddings were L2-normalised. The final integration used gallery embeddings and "
        "distance-threshold matching.",
        size=8.2,
    )
    add_figure_note(document, 1, "metric_kaixiang_roc_curves.png", "show ROC-AUC comparison for my four metric models")
    add_figure_note(document, 2, "metric_best_score_distribution.png", "show same-person vs different-person score separation for the selected Triplet ResNet18 model")

    add_heading(document, "4.2 Anti-Spoofing / Liveness Detection", level=2)
    liveness_rows = [
        ["MobileNetV2 + binary head", "BCE", "0.9447", "0.8750", "0.8756", "0.8713", "0.8800", "62.74"],
        ["EfficientNetB0 + binary head", "BCE", "0.9724", "0.9150", "0.9128", "0.9368", "0.8900", "59.27"],
    ]
    add_table(document, ["Model", "Loss", "AUC", "Acc.", "F1", "Precision", "Recall", "FPS"], liveness_rows)
    add_para(
        document,
        "EfficientNetB0 + binary head was selected as the final liveness model because it achieved the highest "
        "ROC-AUC, accuracy, and F1-score. Although MobileNetV2 remained highly competitive and fast, liveness is "
        "security-critical, so spoof-detection reliability was prioritised.",
        size=8.2,
    )
    add_figure_note(document, 3, "liveness_kaixiang_model_comparison.png", "compare MobileNetV2 and EfficientNetB0 liveness performance")

    add_heading(document, "4.3 Innovative Feature: Glasses Detection", level=2)
    add_para(
        document,
        "My innovative feature was glasses detection. The rationale was to add extra appearance metadata to the "
        "attendance system without interfering with the core identity and liveness decisions. In a real attendance "
        "interface, this feature can display whether the recognised person is wearing glasses, making the system "
        "more informative and demonstrating an additional independent vision task built from the same detected face crop.",
        size=8.3,
    )
    add_para(
        document,
        "The task was formulated as binary classification: with_glasses or without_glasses. I trained MobileNetV2 "
        "and EfficientNetB0 backbones with newly added binary heads. Training followed the same transfer-learning "
        "principle as liveness: first train the binary head while the backbone is frozen, then fine-tune the final "
        "backbone blocks. Class weighting was used because the dataset had more non-glasses samples than glasses samples.",
        size=8.3,
    )
    glasses_rows = [
        ["MobileNetV2 + binary head", "BCEWithLogits", "0.9932", "0.9892", "0.9942", "0.9917", "119.12"],
        ["EfficientNetB0 + binary head", "BCEWithLogits", "0.9923", "0.9884", "0.9927", "0.9906", "103.52"],
    ]
    add_table(document, ["Model", "Loss", "Acc.", "Precision", "Recall", "F1", "FPS"], glasses_rows)
    add_para(
        document,
        "MobileNetV2 + binary head was selected for the glasses feature because it achieved the best F1-score and "
        "higher FPS. This made it suitable for real-time webcam demonstration while still achieving very high accuracy.",
        size=8.2,
    )
    add_figure_note(document, 4, "glasses_kaixiang_model_comparison.png", "summarise glasses model accuracy, F1-score, and FPS")

    add_heading(document, "5. Challenges and Resolutions", level=1)
    add_bullet(document, "Metric recognition instability: similar-looking users sometimes produced very close embeddings in webcam tests. I improved the model using triplet training and used distance thresholds, ambiguity margins, gallery embeddings, and temporal smoothing for integration testing.")
    add_bullet(document, "Dataset limitations: the original verification dataset was not stable enough for my pipeline, so I used VGGFace2 and evaluated on identity-disjoint test identities with fixed verification pairs.")
    add_bullet(document, "Liveness crop sensitivity: tight crops could miss spoof context, so I tested expanded liveness crops and tuned the threshold for stronger real-world anti-spoofing behaviour.")
    add_bullet(document, "Runtime constraints: CPU training was slow, so I used staged fine-tuning, early stopping, and capped glasses samples to keep experiments manageable.")

    add_heading(document, "6. Reflection", level=1)
    add_para(
        document,
        "The final selected models were Triplet ResNet18 for metric recognition, EfficientNetB0 + binary head for "
        "liveness detection, and MobileNetV2 + binary head for glasses detection. The key lesson was that offline "
        "metrics such as ROC-AUC are essential, but final webcam behaviour also depends on crop quality, lighting, "
        "face pose, and gallery quality. Combining quantitative testing with live integration checks made the final "
        "modules more suitable for the attendance system.",
    )

    document.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
